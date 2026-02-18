from __future__ import annotations
from types import SimpleNamespace
from dataclasses import dataclass
from typing import Optional
import pathlib
from docx import Document  # python-docx


@dataclass
class ResumeContent:
    raw_text: str
    source: str  # "docx" | "text"


def extract_text_from_docx(path: str | pathlib.Path) -> str:
    doc = Document(str(path))
    parts = []

    # Extract paragraphs
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    # Extract tables if present
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                (cell.text or "").strip().replace("\n", " ")
                for cell in row.cells
            ).strip()
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


from pathlib import Path

def load_resume(file_path, pasted_text=None):
    path = Path(file_path)

    # If pasted text exists, trust it
    if pasted_text:
        return SimpleNamespace(
            raw_text=pasted_text,
            source="pasted_text",
        )

    # DOCX handling
    if path.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return SimpleNamespace(
            raw_text=text,
            source=path.name,
        )

    # PDF handling
    if path.suffix.lower() == ".pdf":
        import pdfplumber
        pages = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                txt = page.extract_text()
                if txt:
                    pages.append(txt)

        if not pages:
            raise ValueError("PDF parsed but no extractable text found.")

        return SimpleNamespace(
            raw_text="\n".join(pages),
            source=path.name,
        )

    raise ValueError(f"Unsupported resume format: {path.suffix}")

